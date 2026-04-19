from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ===============================
# CONTENT
# ===============================

TITLE = "Smart Attendance System Using Face Recognition"

ABSTRACT = """Attendance is an important part of every educational system, but traditional methods are often slow and unreliable. In this project, a smart attendance system is developed using face recognition. The system captures images through a webcam and identifies students automatically. Once identified, attendance is stored in a database. This reduces manual work and improves accuracy."""

INTRODUCTION = """Manual attendance systems are time-consuming and can be manipulated easily. This project provides an automated solution using face recognition which works through a webcam and removes the need for manual effort."""

PROBLEM = """The main problem is inefficiency and inaccuracy in traditional attendance systems. The goal is to automate attendance using face recognition."""

METHODOLOGY = """The system captures images, detects faces, converts them into encodings, and compares them with stored data to identify students."""

IMPLEMENTATION = """The project is built using Python, Flask, OpenCV, and PostgreSQL. The frontend allows interaction while backend processes recognition and stores data."""

RESULTS = """The system successfully detects and marks attendance. It avoids duplicate entries and works efficiently under normal conditions."""

CONCLUSION = """The system provides a reliable and automated way to manage attendance. It reduces errors and saves time."""

# ===============================
# BUILD DOCUMENT
# ===============================

def build_doc():
    doc = Document()

    # margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_text(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sections
    add_heading("Abstract")
    add_text(ABSTRACT)

    add_heading("Introduction")
    add_text(INTRODUCTION)

    add_heading("Problem Statement")
    add_text(PROBLEM)

    add_heading("Methodology")
    add_text(METHODOLOGY)

    add_heading("Implementation")
    add_text(IMPLEMENTATION)

    add_heading("Results")
    add_text(RESULTS)

    add_heading("Conclusion")
    add_text(CONCLUSION)

    return doc


# ===============================
# MAIN
# ===============================

if __name__ == "__main__":
    print("🔥 STARTING SCRIPT")

    output_path = "Smart_Attendance_System.docx"

    try:
        doc = build_doc()
        doc.save(output_path)

        print("✅ FILE CREATED SUCCESSFULLY")
        print(f"📄 Location: {output_path}")

    except Exception as e:
        print("❌ ERROR:", e)